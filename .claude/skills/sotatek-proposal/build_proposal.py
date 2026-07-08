#!/usr/bin/env python3
"""Generate a Sotatek-branded .docx proposal from a markdown file.

Usage:
    python build_proposal.py INPUT.md [OUTPUT.docx] [--template TEMPLATE.docx]

The template's cover, header/footer branding, styles and numbering are preserved;
only the body content is replaced from the markdown. A real Word TOC field is
inserted (page numbers fill in on open in Word / on the QA render) so no
two-pass rendering is needed.
"""
import argparse
import os
import sys

from docx import Document
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from lib.md_parser import parse
from lib.cover import apply_cover
from lib import docx_helpers as H

REPO = "/Users/trung.hoang/Desktop/presale-sotatek/agency-brands-ads"
DEFAULT_TEMPLATE = os.path.join(REPO, "template", "Sotatek_Naisiti_Project_Proposal.docx")


def _wt(el):
    return "".join(t.text or "" for t in el.iter(H.W("t")))


def _pstyle(el):
    pPr = el.find(qn("w:pPr"))
    if pPr is None:
        return None
    ps = pPr.find(qn("w:pStyle"))
    return ps.get(qn("w:val")) if ps is not None else None


def _has_toc_instr(el):
    """True if the element wraps a TOC field (the 'TOC' switch is in instrText,
    not in visible <w:t>, so a text-only check misses it)."""
    for it in el.iter(qn("w:instrText")):
        if it.text and "TOC" in it.text:
            return True
    return False


def clear_body_and_toc(doc):
    """Remove stale sdt-TOC controls, empty front-matter tails and the old body
    (everything from the first Heading 1 onward). Return the TOC-title element."""
    body = doc.element.body

    # 1) drop stale TOC content-controls (Google-Docs export wraps TOC in <w:sdt>)
    for ch in list(body):
        if ch.tag == qn("w:sdt") and _has_toc_instr(ch):
            body.remove(ch)

    # 2) find the "Table of Contents" title paragraph
    toc_title = None
    for ch in body.iterchildren():
        if ch.tag == qn("w:p") and "Table of Contents" in _wt(ch):
            toc_title = ch
            break

    # 3) remove the old body: first Heading1 paragraph and all following p/tbl
    kids = list(body.iterchildren())
    start = None
    for i, ch in enumerate(kids):
        if ch.tag == qn("w:p") and _pstyle(ch) == "Heading1":
            start = i
            break
    if start is not None:
        for ch in kids[start:]:
            if ch.tag in (qn("w:p"), qn("w:tbl")):
                body.remove(ch)

    # 4) remove every empty paragraph between the TOC title and the section prop
    #    (leftover blank lines / empty headings that would pollute the TOC)
    if toc_title is not None:
        sib = toc_title.getnext()
        while sib is not None and sib.tag != qn("w:sectPr"):
            nxt = sib.getnext()
            if sib.tag == qn("w:p") and not _wt(sib).strip():
                body.remove(sib)
            sib = nxt
    return toc_title


def emit_blocks(doc, blocks, headings):
    first_heading = True
    for b in blocks:
        t = b["type"]
        if t == "heading":
            H.add_heading(doc, b["text"], b["level"], headings=headings,
                          page_break=(first_heading and b["level"] == 1))
            if b["level"] == 1:
                first_heading = False
        elif t == "para":
            H.add_para(doc, b["text"])
        elif t == "bullet":
            H.add_bullet(doc, b["text"], level=b["level"])
        elif t == "note":
            H.add_para(doc, b["text"], italic=True, size=9)
        elif t == "table":
            H.add_table(doc, b["rows"], align=b.get("align"))


def strip_unused_media(path):
    """Remove media parts (e.g. the template's Naisiti architecture diagram) whose
    relationship id is no longer referenced in document.xml, plus their rel entry."""
    import re
    import zipfile
    z = zipfile.ZipFile(path)
    rels = z.read("word/_rels/document.xml.rels").decode()
    docx = z.read("word/document.xml").decode()
    drop_files, drop_ids = set(), set()
    for m in re.finditer(r'<Relationship\b[^>]*/>', rels):
        rel = m.group(0)
        rid = re.search(r'Id="([^"]+)"', rel)
        tgt = re.search(r'Target="(media/[^"]+)"', rel)
        if rid and tgt and ('r:embed="%s"' % rid.group(1)) not in docx \
                and ('r:id="%s"' % rid.group(1)) not in docx:
            drop_ids.add(rid.group(1))
            drop_files.add("word/" + tgt.group(1))
    if not drop_ids:
        z.close()
        return 0
    tmp = path + ".tmp"
    zout = zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED)
    for item in z.infolist():
        if item.filename in drop_files:
            continue
        data = z.read(item.filename)
        if item.filename == "word/_rels/document.xml.rels":
            for rid in drop_ids:
                data = re.sub(
                    r'<Relationship\b[^>]*Id="%s"[^>]*/>' % re.escape(rid),
                    "", data.decode()).encode()
        zout.writestr(item, data)
    z.close(); zout.close()
    os.replace(tmp, path)
    return len(drop_files)


def _assemble(fm, blocks, template_path, toc_mode, pages):
    """Build the doc once. toc_mode: 'field' (no numbers) or 'static' (numbers from
    `pages`). Returns (doc, headings)."""
    doc = Document(template_path)
    apply_cover(doc, fm)
    toc_title = clear_body_and_toc(doc)

    headings = []
    emit_blocks(doc, blocks, headings)

    if toc_title is not None:
        if toc_mode == "field":
            toc_title.addnext(H.make_toc_field(doc))
            H.set_update_fields(doc)
        else:  # static: one entry per heading, page numbers from `pages`
            anchor = toc_title
            for level, text in headings:
                el = H.make_toc_entry(doc, level, text,
                                      (pages or {}).get(text.strip(), ""))
                anchor.addnext(el)
                anchor = el
    return doc, headings


def build(md_path, out_path, template_path, fast=False):
    text = open(md_path, encoding="utf-8").read()
    fm, blocks = parse(text)

    if fast:
        doc, headings = _assemble(fm, blocks, template_path, "field", None)
        doc.save(out_path)
        strip_unused_media(out_path)
        return len(headings), len(doc.tables), "field"

    # default: static populated TOC via one render + rebuild with real page numbers
    from lib.render_map import render_pdf, map_heading_pages
    doc, headings = _assemble(fm, blocks, template_path, "static", None)
    doc.save(out_path)                       # placeholder page numbers
    try:
        pdf = render_pdf(out_path)
        pages = map_heading_pages(pdf, headings)
        doc, headings = _assemble(fm, blocks, template_path, "static", pages)
        doc.save(out_path)                   # correct page numbers (same layout)
        mode = "static"
    except Exception as e:  # never hard-fail: fall back to a field TOC
        doc, headings = _assemble(fm, blocks, template_path, "field", None)
        doc.save(out_path)
        print("WARN: static TOC render failed (%s); wrote field TOC instead." % e)
        mode = "field (fallback)"
    strip_unused_media(out_path)
    return len(headings), len(doc.tables), mode


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input")
    ap.add_argument("output", nargs="?")
    ap.add_argument("--template", default=DEFAULT_TEMPLATE)
    ap.add_argument("--fast", action="store_true",
                    help="field TOC, no render (populates on open in Word)")
    a = ap.parse_args()
    out = a.output or os.path.join(
        REPO, "output",
        "Sotatek_" + os.path.splitext(os.path.basename(a.input))[0] + ".docx")
    nh, nt, mode = build(a.input, out, a.template, fast=a.fast)
    print("SAVED:", out)
    print("headings:", nh, "| tables:", nt, "| TOC:", mode)


if __name__ == "__main__":
    main()
