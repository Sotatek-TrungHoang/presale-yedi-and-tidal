"""Low-level docx helpers matching the Sotatek template's look.

Brand: Heading1 navy #1A2B4A, Heading2 blue #2E75B6, Heading3 orange #D4740E,
Arial body, bullets via numId 5 (defined in the template's numbering.xml).
"""
import re
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = "1A2B4A"
GREY = "D9D9D9"
LGREY = "F2F5F9"
BODY = "222222"
BULLET_NUMID = "5"          # bullet list defined in template numbering.xml
CONTENT_W = 9360            # usable page width (twips)
TOC_TAB = 9180              # TOC page-number tab: just inside the right margin so
                            # the right-aligned number always renders (a tab exactly
                            # at the edge is dropped by some renderers)

W = lambda tag: qn("w:" + tag)


# ---- run / text ----
def style_run(run, size=None, bold=None, italic=None, color=None, font="Arial"):
    run.font.name = font
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_richtext(p, text, size=11, bold=False, italic=False, color=BODY):
    """Add runs to paragraph, honouring **bold** and *italic* inline markers."""
    for tok in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text):
        if not tok:
            continue
        b, i, t = bold, italic, tok
        if tok.startswith("**") and tok.endswith("**"):
            b, t = True, tok[2:-2]
        elif tok.startswith("*") and tok.endswith("*"):
            i, t = True, tok[1:-1]
        style_run(p.add_run(t), size=size, bold=b, italic=i, color=color)


def set_para_wt(el, text):
    """Replace the first <w:t> descendant of a paragraph element, blank the rest
    (preserves the original run styling; handles hyperlink-wrapped runs)."""
    ts = [e for e in el.iter(W("t"))]
    if ts:
        ts[0].text = text
        for e in ts[1:]:
            e.text = ""


def reset_cell(cell, text, bold=False, size=10, color=BODY, align=None):
    """Wipe a cell's first paragraph (runs AND hyperlinks) and write one clean run."""
    p = cell.paragraphs[0]
    for child in list(p._p):
        if child.tag in (W("r"), W("hyperlink")):
            p._p.remove(child)
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, color=color)
    if align:
        p.alignment = {"l": WD_ALIGN_PARAGRAPH.LEFT, "c": WD_ALIGN_PARAGRAPH.CENTER,
                       "r": WD_ALIGN_PARAGRAPH.RIGHT}[align]


# ---- shading / borders ----
def _shade(parent, fill):
    shd = OxmlElement("w:shd")
    shd.set(W("val"), "clear"); shd.set(W("color"), "auto"); shd.set(W("fill"), fill)
    parent.append(shd)


def cell_bg(cell, fill):
    _shade(cell._tc.get_or_add_tcPr(), fill)


def _table_borders(tbl):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(W("val"), "single"); e.set(W("sz"), "4")
        e.set(W("space"), "0"); e.set(W("color"), GREY)
        borders.append(e)
    tbl._tbl.tblPr.append(borders)


def _col_widths(rows):
    ncol = len(rows[0])
    ml = [1] * ncol
    for r in rows:
        for i in range(ncol):
            ml[i] = max(ml[i], len(str(r[i])) if i < len(r) else 1)
    s = sum(ml)
    w = [max(700, int(CONTENT_W * m / s)) for m in ml]
    f = CONTENT_W / sum(w)
    return [int(x * f) for x in w]


# ---- block builders (append to end of body, before sectPr) ----
def add_heading(doc, text, level, headings=None, page_break=False):
    p = doc.add_paragraph(style="Heading %d" % level)
    if page_break:
        p.paragraph_format.page_break_before = True
    p.add_run(text)
    if headings is not None:
        headings.append((level, text))
    return p


def add_para(doc, text, italic=False, size=11, space_after=6):
    p = doc.add_paragraph(style="normal")
    add_richtext(p, text, size=size, italic=italic,
                 color=("555555" if italic else BODY))
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="normal")
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    il = OxmlElement("w:ilvl"); il.set(W("val"), str(level)); numPr.append(il)
    ni = OxmlElement("w:numId"); ni.set(W("val"), BULLET_NUMID); numPr.append(ni)
    pPr.append(numPr)
    ind = OxmlElement("w:ind")
    ind.set(W("left"), str(720 + level * 360)); ind.set(W("hanging"), "360")
    pPr.append(ind)
    p.paragraph_format.space_after = Pt(3)
    add_richtext(p, text, size=11)
    return p


def add_table(doc, rows, align=None, header=True, fontsize=10):
    ncol = len(rows[0])
    align = align or ["l"] * ncol
    if len(align) < ncol:
        align += ["l"] * (ncol - len(align))
    widths = _col_widths(rows)
    tbl = doc.add_table(rows=len(rows), cols=ncol)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_borders(tbl)
    amap = {"l": WD_ALIGN_PARAGRAPH.LEFT, "c": WD_ALIGN_PARAGRAPH.CENTER,
            "r": WD_ALIGN_PARAGRAPH.RIGHT}
    for ri, row in enumerate(rows):
        for ci in range(ncol):
            val = row[ci] if ci < len(row) else ""
            cell = tbl.cell(ri, ci)
            para = cell.paragraphs[0]
            para.alignment = amap[align[ci]]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            head = header and ri == 0
            add_richtext(para, str(val), size=fontsize,
                         bold=head, color=("FFFFFF" if head else BODY))
            if head:
                cell_bg(cell, NAVY)
            elif ri % 2 == 0:
                cell_bg(cell, LGREY)
            cell.width = Twips(widths[ci])
    for row in tbl.rows:
        for ci, c in enumerate(row.cells):
            c.width = Twips(widths[ci])
    return tbl


# ---- Table of Contents field ----
def make_toc_field(doc):
    """Return a paragraph element holding a real Word TOC field (levels 1-3)."""
    p = doc.add_paragraph(style="normal")
    r = p.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(W("fldCharType"), "begin"); r._r.append(fb)
    r2 = p.add_run()
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = ' TOC \\o "1-3" \\h \\z \\u '; r2._r.append(it)
    r3 = p.add_run()
    fs = OxmlElement("w:fldChar"); fs.set(W("fldCharType"), "separate"); r3._r.append(fs)
    r4 = p.add_run("Update this field to generate the table of contents "
                   "(select all, then press F9 in Word).")
    style_run(r4, size=10, italic=True, color="777777")
    r5 = p.add_run()
    fe = OxmlElement("w:fldChar"); fe.set(W("fldCharType"), "end"); r5._r.append(fe)
    el = p._p
    el.getparent().remove(el)
    return el


def set_update_fields(doc):
    """Make Word refresh fields (TOC page numbers) on first open."""
    settings = doc.settings.element
    if settings.find(W("updateFields")) is None:
        uf = OxmlElement("w:updateFields"); uf.set(W("val"), "true")
        settings.append(uf)


def make_toc_entry(doc, level, text, page):
    """Return a detached paragraph element: one static TOC line with a right-aligned
    dot-leader tab to the page number. Populated everywhere (no Word update needed)."""
    p = doc.add_paragraph(style="normal")
    pf = p.paragraph_format
    pf.left_indent = Twips({1: 0, 2: 300, 3: 620}.get(level, 620))
    pf.space_after = Pt(2); pf.space_before = Pt(0)
    pf.tab_stops.add_tab_stop(Twips(TOC_TAB), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    size = 11 if level == 1 else 10
    color = NAVY if level == 1 else "333333"
    style_run(p.add_run(text), size=size, bold=(level == 1), color=color)
    style_run(p.add_run("\t%s" % (page if page else "")),
              size=size, bold=(level == 1), color="333333")
    el = p._p
    el.getparent().remove(el)
    return el
